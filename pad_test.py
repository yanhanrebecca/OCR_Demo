import cv2
import numpy as np
import subprocess
import os
import argparse
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage

# 初始化参数
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'     # 环境变量设置：设置 KMP_DUPLICATE_LIB_OK 可以避免程序因库冲突而崩溃
HIOG = 5 # 水平投影的阈值 30 3；5 8；
VIOG = 5   # 垂直投影的阈值
Position = []

'''计算二值图像的水平投影'''
def getHProjection(image):
    # 创建一个与输入图像image相同形状的空白图像（用零填充），用于可视化水平投影。
    hProjection = np.zeros(image.shape, np.uint8)
    # 获取图像大小
    (h, w) = image.shape
    # 统计像素个数
    h_ = [0] * h
    #计算水平投影：如果像素为白色（值为255），则在 h_[y] 对应的位置增加计数。
    for y in range(h):
        for x in range(w):
            if image[y, x] == 255:
                h_[y] += 1
    # 绘制水平投影图像
    for y in range(h):
        for x in range(h_[y]):
            hProjection[y, x] = 255
    return h_

'''计算二值化图像 image 的垂直投影'''
def getVProjection(image):
    vProjection = np.zeros(image.shape, np.uint8)
    (h, w) = image.shape
    w_ = [0] * w
    for x in range(w):
        for y in range(h):
            if image[y, x] == 255:
                w_[x] += 1
    for x in range(w):
        for y in range(h - w_[x], h):
            vProjection[y, x] = 255
    return w_

'''扫描垂直投影 vProjection，并找到投影中连续超过给定阈值 iog 的区间'''
def scan(projection, threshold, min_distance=0):
    start = 0       # 用于指示当前是否在一个有效的区间内。初始化为 0，表示当前不在有效区间内
    starts = []     # 用于存储有效区间的起始位置。
    ends = []       # 用于存储有效区间的结束位置
    for i in range(len(projection)):
        # 进入区间
        if projection[i] > threshold and start == 0:
            starts.append(i)
            start = 1
        #退出区间
        if projection[i] <= threshold and start == 1:
            if i - starts[-1] < min_distance:
                continue
            ends.append(i)
            start = 0
    return starts, ends


# 裁剪图像并保存裁剪后的图像
def CropImage(image, dest, boxMin, boxMax):     #dest：裁剪后图像的保存路径和文件名。boxMin：裁剪框的左上角坐标 (x_min, y_min)，形式为 (行, 列)
    try:
        # Extract cropping bounds
        a = boxMin[1]       # a 和 b 分别是裁剪框的纵向边界（即图像的行范围）
        b = boxMax[1]
        c = boxMin[0]       # c 和 d 分别是裁剪框的横向边界（即图像的列范围）
        d = boxMax[0]

        # Perform image cropping
        cropImg = image[a:b, c:d]

        # Check if cropImg is valid (not empty)
        if cropImg.size == 0:
            print(f"Warning: 裁剪图像为空. BoxMin: {boxMin}, BoxMax: {boxMax}")
            return

        # 增强裁剪后的图像
        enhancedImg = enhance_image(cropImg)

        # 保存增强后的图像
        cv2.imwrite(dest, enhancedImg)
    # 异常处理
    except Exception as e:
        print(f"处理图像时发生错误: {e}")

def enhance_image(image):
    # 将输入的彩色图像转换为灰度图像
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # 使用高斯模糊对灰度图像进行平滑处理，减少图像噪声，参数(5,5)表示高斯核的大小，0表示sigma由核大小自动计算
    smoothed = cv2.GaussianBlur(gray_image, (5, 5), 0)
    # 使用Otsu's二值化方法对平滑后的图像进行二值化处理，自动确定阈值，将图像转换为二值图像
    _, binary = cv2.threshold(smoothed, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    # 计算二值图像的拉普拉斯算子，用于边缘检测，这里先转换为64位浮点数类型，以便处理可能产生的负值
    laplacian = cv2.Laplacian(binary, cv2.CV_64F)
    # 将拉普拉斯算子的结果转换回8位无符号整数类型，并取绝对值，因为拉普拉斯算子可能产生负值
    laplacian = cv2.convertScaleAbs(laplacian)
    # 将二值图像和拉普拉斯算子的结果（边缘信息）进行加权合并，以增强边缘效果
    # 参数1.5和0.5分别是二值图像和拉普拉斯算子的权重，最后一个参数0是加权的和
    edge_enhanced = cv2.addWeighted(binary, 1.5, laplacian, 0.5, 0)
    return edge_enhanced    # 返回边缘增强的图像

# 开始识别
def DOIT(rawPic, save_directory):
    # 读入原始图像
    origineImage = cv2.imread(rawPic)
    # 图像灰度化
    image = cv2.cvtColor(origineImage, cv2.COLOR_BGR2GRAY)

    # 将图片二值化 原始阈值127
    retval, img = cv2.threshold(image, 127, 255, cv2.THRESH_BINARY_INV)
    (h, w) = img.shape
    # 垂直投影
    V = getVProjection(img)

    start = 0
    V_start = []
    V_end = []

    # 对垂直投影水平分割
    V_start, V_end = scan(V, HIOG)
    if len(V_start) > len(V_end):
        V_end.append(w - 5)

    # 分割行，分割之后再进行列分割并保存分割位置
    for i in range(len(V_end)):
        # 获取行图像 30
        if V_end[i] - V_start[i] < 30:
            continue

        cropImg = img[0:h, V_start[i]:V_end[i]]
        H = getHProjection(cropImg)
        H_start, H_end = scan(H, VIOG, 40)

        if len(H_start) > len(H_end):
            H_end.append(h - 5)

        for pos in range(len(H_start)):
            # 再进行一次列扫描
            DcropImg = cropImg[H_start[pos]:H_end[pos], 0:w]
            d_h, d_w = DcropImg.shape
            sec_V = getVProjection(DcropImg)
            c1, c2 = scan(sec_V, 0)
            if len(c1) > len(c2):
                c2.append(d_w)

            x = 1
            while x < len(c1):
                if c1[x] - c2[x - 1] < 12:
                    c2.pop(x - 1)
                    c1.pop(x)
                    x -= 1
                x += 1

            if len(c1) == 1:
                Position.append([V_start[i], H_start[pos], V_end[i], H_end[pos]])
            else:
                for x in range(len(c1)):
                    Position.append([V_start[i] + c1[x], H_start[pos], V_start[i] + c2[x], H_end[pos]])

    # 根据确定的位置分割字符
    number = 0
    for m in range(len(Position)):
        rectMin = (Position[m][0], Position[m][1])
        rectMax = (Position[m][2], Position[m][3])
        sizes = [(0, 0), (50, 50), (100, 100)]  # 定义三种不同的尺寸

        for size in sizes:
            # 调整裁剪区域以确保不引入黑色边框
            newRectMin = (rectMin[0] - size[0], rectMin[1] - size[1])
            newRectMax = (rectMax[0] + size[0], rectMax[1] + size[1])

            # 保存裁剪后的图像到指定目录
            crop_path = os.path.join(save_directory, f'{number}.jpg')
            CropImage(origineImage, crop_path, newRectMin, newRectMax)
            number += 1

    result_image_path = os.path.join(save_directory, 'ResultImage.jpg')
    cv2.imwrite(result_image_path, origineImage)



def ocr(image_dir):
    picture_dir = os.path.join(image_dir, 'picture')
    unrecognized_dir = os.path.join(image_dir, 'unrecognized')

    if not os.path.exists(picture_dir):
        os.makedirs(picture_dir)
    if not os.path.exists(unrecognized_dir):
        os.makedirs(unrecognized_dir)

    files = os.listdir(image_dir)
    for file in files:
        file_path = os.path.join(image_dir, file)
        if file.lower().endswith('.jpg'):
            print("-----------------------------------------------------------")
            print("正在识别：{}".format(file))
            pad_ocr(file_path, picture_dir, unrecognized_dir)

    # 创建一个新的 Word 文档
    doc = Document()

    # 遍历 picture 目录中的每个子目录
    for foldername, subfolders, filenames in os.walk(picture_dir):
        if filenames:
            # 只处理图片文件
            image_files = [f for f in filenames if f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif'))]
            if image_files:
                # 在 Word 文档中插入一行
                row_cells = doc.add_table(rows=1, cols=len(image_files)).rows[0].cells
                for i, filename in enumerate(image_files):
                    image_path = os.path.join(foldername, filename)

                    # 检查图片格式是否被支持
                    try:
                        with PILImage.open(image_path) as img:
                            img.verify()  # 确保图片格式正确
                    except (IOError, SyntaxError) as e:
                        print(f"警告: 无法识别图片文件 {image_path}. 错误信息: {e}")
                        continue

                    # 添加图片到 Word 文档
                    try:
                        cell = row_cells[i]
                        run = cell.add_paragraph().add_run()
                        run.add_picture(image_path, width=Inches(1.0))  # 调整图片宽度

                        # 添加图片说明
                        para = cell.add_paragraph()
                        para.add_run(f"来自 {filename}")
                    except Exception as e:
                        print(f"警告: 无法将图片 {image_path} 插入到 Word 文档. 错误信息: {e}")
                        continue

    # 保存 Word 文档到 picture 目录
    doc_path = os.path.join(picture_dir, 'output.docx')
    doc.save(doc_path)
    print(f"Word 文档已保存至：{doc_path}")


# 执行 OCR（光学字符识别）操作，处理图像文件，并将识别结果进行保存和分类
def pad_ocr(file_path, picture_dir, unrecognized_dir):      #file_path：待处理图像文件的完整路径
    # 命令行调用 paddleocr 工具的命令，构建命令
    command = ['paddleocr', '--image_dir', file_path, '--use_angle_cls', 'true', '--use_gpu', 'true']
    # 执行命令并捕获输出，使用 subprocess.run 执行构建的命令   capture_output=True：捕获命令的标准输出和标准错误输出    text=True：将输出处理为字符串而不是字节
    result = subprocess.run(command, capture_output=True, text=True)

    # 检查命令是否成功执行
    if result.returncode == 0:
        print("识别成功。")
        output = result.stdout  # 获取命令的标准输出
        if output:
            output_lines = output.splitlines()  # 将输出按行分割
            if output_lines:
                last_line = output_lines[-1]    # 获取输出的最后一行，这通常是 OCR 识别的结果
                word = process_log(last_line)   # 处理这行文本以提取识别出的文字
                word = word.replace("'", "")    # 去除识别结果中的单引号（如果有的话）
                print("{}图片是识别出的文字是{}".format(file_path, word))
                if word:
                    word_folder = os.path.join(picture_dir, word)
                    if not os.path.exists(word_folder):
                        os.makedirs(word_folder)
                    txt_path = os.path.join(word_folder, f'{word}.txt')
                    with open(txt_path, 'w') as f:
                        f.write(word)
                    move_file(file_path, word_folder)
                else:
                    move_file(file_path, unrecognized_dir)
            else:
                print("命令没有输出任何内容。")
                move_file(file_path, unrecognized_dir)
        else:
            print("标准输出是 None。")
            move_file(file_path, unrecognized_dir)
    else:
        print("命令执行失败。")
        # 打印或处理命令的错误输出
        print("错误输出:")
        print(result.stderr)
        move_file(file_path, unrecognized_dir)


#将文件移动到指定的目标目录
def move_file(file_path, target_dir, word=None):
    # 处理子目录
    if word:
        target_dir = os.path.join(target_dir, word)
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)
    base_name = os.path.basename(file_path)
    target_path = os.path.join(target_dir, base_name)
    # 生成新的文件名（如 name_1.ext）
    if os.path.exists(target_path):
        name, ext = os.path.splitext(base_name)
        counter = 1
        while os.path.exists(target_path):
            new_name = f"{name}_{counter}{ext}"
            target_path = os.path.join(target_dir, new_name)
            counter += 1
    os.rename(file_path, target_path)   # 将源文件从 file_path 移动到目标路径 target_path

#处理字符串 log_string，从中提取特定的信息
def process_log(log_string):
    start_index = log_string.find('(')
    if start_index != -1:
        # 截取左括号后面的内容
        content_after_bracket = log_string[start_index + 1:]
        # 找到逗号的位置，即找到左括号后面的内容的结束位置
        end_index = content_after_bracket.find(',')
        if end_index != -1:
            # 提取找到的字符
            found_character = content_after_bracket[:end_index].strip()
            return found_character  # 提取从 content_after_bracket 的开始到逗号位置 end_index 的所有字符，返回提取的内容
    return ""   # 未找到括号或逗号的处理

# 主程序部分：处理传入的图片，执行一些图像处理操作，并在处理完毕后进行 OCR
if __name__ == '__main__':
    # 设置保存目录为D盘某个路径
    save_directory = r'./result'
    # 确保保存目录存在
    if not os.path.exists(save_directory):
        os.makedirs(save_directory)
    # 示例
    # rawPicPath = r"./2.jpg"
    print("开始图片处理")
    # 对命令行处理
    parser = argparse.ArgumentParser()
    parser.add_argument('-img', '--image_path', help='输入图片路径')
    args = parser.parse_args()
    # 获取传入的参数值
    rawPicPath = args.image_path
    # 调用 DOIT 函数进行图片处理
    DOIT(rawPicPath, save_directory)
    print("图片处理完毕")
    image_dir = './result'
    print("传入的图片路径是:", rawPicPath)
    # print(args)
    # 调用 ocr 函数进行 OCR 处理
    ocr(image_dir)